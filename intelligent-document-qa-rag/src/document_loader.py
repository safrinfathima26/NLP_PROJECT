"""
document_loader.py
-------------------
Handles loading raw text out of user-supplied documents.

Supported formats: .pdf, .docx, .txt

Each loader function returns a single string containing the
full extracted text of the document. Page/paragraph boundaries
are preserved with newlines so downstream chunking can make
reasonable splitting decisions.
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Union

from pypdf import PdfReader
from docx import Document as DocxDocument


class DocumentLoadError(Exception):
    """Raised when a document cannot be read or parsed."""


def load_pdf(file_path: Union[str, Path]) -> str:
    """Extract text from a PDF file, page by page."""
    reader = PdfReader(str(file_path))
    pages_text = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages_text.append(f"[Page {page_num}]\n{text.strip()}")
    if not pages_text:
        raise DocumentLoadError(
            f"No extractable text found in PDF: {file_path}. "
            "The file may be a scanned image without OCR."
        )
    return "\n\n".join(pages_text)


def load_docx(file_path: Union[str, Path]) -> str:
    """Extract text from a Word (.docx) file, paragraph by paragraph."""
    doc = DocxDocument(str(file_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Also pull text out of tables, since many reports keep key data there
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    if not paragraphs:
        raise DocumentLoadError(f"No extractable text found in DOCX: {file_path}")
    return "\n".join(paragraphs)


def load_txt(file_path: Union[str, Path]) -> str:
    """Load a plain text file, trying a couple of common encodings."""
    path = Path(file_path)
    for encoding in ("utf-8", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentLoadError(f"Could not decode text file: {file_path}")


def load_document(file_path: Union[str, Path]) -> str:
    """
    Dispatch to the correct loader based on file extension.

    Parameters
    ----------
    file_path : str | Path
        Path to a .pdf, .docx, or .txt file.

    Returns
    -------
    str
        The extracted plain text of the document.
    """
    path = Path(file_path)
    if not path.exists():
        raise DocumentLoadError(f"File not found: {file_path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return load_pdf(path)
    elif suffix == ".docx":
        return load_docx(path)
    elif suffix == ".txt":
        return load_txt(path)
    else:
        raise DocumentLoadError(
            f"Unsupported file type '{suffix}'. Supported types: .pdf, .docx, .txt"
        )


def load_documents_from_directory(directory: Union[str, Path]) -> dict:
    """
    Load every supported document in a directory.

    Returns
    -------
    dict[str, str]
        Mapping of {filename: extracted_text} for every file that
        was successfully loaded. Files that fail to load are skipped
        with a printed warning rather than crashing the whole batch.
    """
    directory = Path(directory)
    supported = {".pdf", ".docx", ".txt"}
    results = {}

    for file_path in sorted(directory.iterdir()):
        if file_path.suffix.lower() in supported:
            try:
                results[file_path.name] = load_document(file_path)
            except DocumentLoadError as e:
                print(f"[WARN] Skipping {file_path.name}: {e}")
    return results
